import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(30, 58, 138)
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(37, 99, 235)
    return p

def create_procedure_docx(output_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    
    title_run = title_p.add_run("PROCEDIMIENTO OPERATIVO ESTANDARIZADO (POE)")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title_run.bold = True
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(14)
    sub_run = subtitle_p.add_run("Mantenimiento Preventivo, Correctivo, Herrería, Controles de Fluidos y Auxilios en Ruta\nContenedores Hugo S.A. — Bajo Norma ISO 9001:2015")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(71, 85, 105)
    sub_run.italic = True

    # Metadata Header Table
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Código:", True), ("POE-MANT-01", False), ("Versión:", True), ("3.2 (Chequeador & Auxilios)", False)],
        [("Vigencia:", True), ("Julio 2026", False), ("Norma ISO:", True), ("9001:2015 (Cl. 7, 8, 9, 10)", False)]
    ]
    for r_idx, row in enumerate(meta_table.rows):
        for c_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "F1F5F9")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            text, is_bold = meta_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.bold = is_bold
            run.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Objetivos y Alcance
    add_heading_styled(doc, "1. Objetivos y Alcance", 1)
    
    add_heading_styled(doc, "1.1 Objetivo", 2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Garantizar la disponibilidad operativa, confiabilidad y seguridad de la flota de la empresa mediante un sistema estructurado de mantenimiento preventivo, correctivo, de herrería, controles periódicos de fluidos y respuesta rápida ante auxilios mecánicos en ruta, asegurando la trazabilidad total bajo la norma ISO 9001:2015.")
    r.font.name = 'Calibri'; r.font.size = Pt(10.5)

    add_heading_styled(doc, "1.2 Alcance", 2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Aplica a todas las unidades de la empresa (Camiones, Camionetas, Autos, Auto-elevadores y Maquinarias) y a todo el personal involucrado (Conductores, Choferes, Chequeadores/Inspectores de Fluidos, Mecánicos, Supervisores, Jefes de Taller, Asistente IA y Auditores Internos).")
    r.font.name = 'Calibri'; r.font.size = Pt(10.5)

    # 2. Mapa de Ecosistema Tecnológico y Enlaces Directos
    add_heading_styled(doc, "2. Mapa de Ecosistema Tecnológico y Enlaces Directos", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("El sistema de gestión de mantenimiento opera sobre una arquitectura distribuida e integrada:")
    r.font.name = 'Calibri'; r.font.size = Pt(10.5)

    eco_rows = [
        ["App Nativa de Gestión (Railway)", "Creación de OT (Preventivo, Correctivo, Herrería, Auxilio), registro de insumos, tiempos, estado de unidades y Hugo AI.", "https://app-taxes-production-ec67.up.railway.app", "Supervisores, Mecánicos, Chequeadores, Jefes"],
        ["Portal Oficial Taxes", "Registro legal y fiscal obligatorio de Órdenes de Trabajo (OT).", "https://taxes.com.ar", "Sincronizador Automático (ASISTENTE IA - Puppeteer Worker)"],
        ["Chatbot WhatsApp / Hugo AI", "Registro instantáneo de novedades, chequeo de fluidos, reporte urgente de auxilios en ruta y consultas AI.", "Módulo Integrado WhatsApp / Chat Hugo AI", "Choferes, Chequeadores, Mecánicos, Supervisores"],
        ["Google Sheets 2026 (Gestión Mantenimiento)", "Base de datos analítica 2026 para ranking de fallas, registro de auxilios en ruta, consumo de combustible e historial.", "https://docs.google.com/spreadsheets/d/1QK698StrEr9v7HgJUrtN1GFtb3ixk_2ql78dkx1_3Vk", "Gerencia, Jefatura de Taller, Auditores, Supervisores"]
    ]

    eco_table = doc.add_table(rows=len(eco_rows) + 1, cols=4)
    eco_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Componente / App", "Función Principal", "Enlace / URL Acceso", "Usuarios Principales"]
    for c_idx, text in enumerate(headers):
        cell = eco_table.rows[0].cells[c_idx]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=90, bottom=90, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row_data in enumerate(eco_rows):
        row_cells = eco_table.rows[r_idx + 1].cells
        bg_hex = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'; r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True
            if c_idx == 2 and val.startswith("http"):
                r.font.color.rgb = RGBColor(37, 99, 235)
                r.underline = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 3. Matriz ISO 9001:2015
    add_heading_styled(doc, "3. Matriz de Cumplimiento ISO 9001:2015", 1)
    
    iso_data = [
        ["7.1.3 Infraestructura", "Mantenimiento de equipos y transporte para asegurar la conformidad del servicio.", "Plan de mantenimiento preventivo y Control Diario de Fluidos por Chequeador (Aceite, Refrigerante, Frenos) según Kilómetros / Litros / Horas."],
        ["8.1 Planificación y Control Operacional", "Establecer criterios para los procesos y ejecución de tareas según especificaciones.", "Creación obligatoria de OT en la App Nativa con clasificación (Preventivo / Correctivo / Herrería / Auxilio en Ruta) y asignación inmediata."],
        ["8.5.1 Control de la Provisión del Servicio", "Control de actividades mediante información documentada e indicadores en tiempo real.", "Registro de tiempo real (cronómetro), insumos de auxilio, reposición de fluidos y validación por Supervisor previo al alta."],
        ["8.7 Control de Salidas No Conformes", "Identificación y aislamiento de equipos defectuosos.", "Cambio automático de estado de unidad a ❌ Fuera de Servicio o 🔧 En Reparación / Auxilio vía Chatbot/App al detectar fallas o fluidos bajos."],
        ["9.1.1 Seguimiento, Medición y Análisis", "Evaluación del desempeño y la eficacia del sistema.", "Sincronización a Google Sheets para análisis de KPI de Auxilios (Frecuencia, Tiempo de Respuesta MTTR, causa raíz)."],
        ["10.2 Mejora Continua", "Corrección de causas raíz y optimización de procesos.", "Auditoría de auxilios recurrentes y aprendizaje continuo con Hugo AI Assistant."]
    ]

    iso_table = doc.add_table(rows=len(iso_data) + 1, cols=3)
    iso_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    iso_headers = ["Cláusula ISO 9001:2015", "Requisito de la Norma", "Aplicación Práctica en Contenedores Hugo"]
    for c_idx, text in enumerate(iso_headers):
        cell = iso_table.rows[0].cells[c_idx]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=90, bottom=90, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row_data in enumerate(iso_data):
        row_cells = iso_table.rows[r_idx + 1].cells
        bg_hex = "F1F5F9" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'; r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 4. Flujograma y Mapa del Proceso VISUAL IN WORD
    add_heading_styled(doc, "4. Mapa del Proceso y Flujograma Operativo", 1)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("4.1 Mapa de Proceso (Macroproceso)")
    r.font.name = 'Calibri'; r.font.size = Pt(11.5); r.bold = True; r.font.color.rgb = RGBColor(30, 58, 138)

    map_table = doc.add_table(rows=2, cols=3)
    map_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    map_headers = ["ENTRADAS", "PROCESO PRINCIPAL", "SALIDAS & ANÁLISIS"]
    for c_idx, text in enumerate(map_headers):
        cell = map_table.rows[0].cells[c_idx]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)

    map_contents = [
        "• Chofer / Novedad WhatsApp\n• Control Diario Fluidos (Chequeador)\n• Auxilio Mecánico en Ruta\n• Plan Preventivos / Herrería",
        "1. Registro & Clasificación (App Nativa)\n2. Despacho Auxilio / Taller\n3. Ejecución & Cronómetro (Mecánicos)\n4. Control & Validación (Supervisor)\n5. Sincronización Oficial (Asistente IA)",
        "• Unidad Operativa en Ruta ✅\n• Google Sheets 2026 (Pestaña Auxilios & KPI)\n• Asistente Hugo AI (Consultas)"
    ]

    for c_idx, text in enumerate(map_contents):
        cell = map_table.rows[1].cells[c_idx]
        set_cell_background(cell, "EFF6FF")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("4.2 Flujograma Operativo Detallado (Incluye Chequeador y Auxilios)")
    r.font.name = 'Calibri'; r.font.size = Pt(11.5); r.bold = True; r.font.color.rgb = RGBColor(30, 58, 138)

    flow_steps = [
        ("INICIO", "Detección de Novedad, Fluidos Faltantes, Preventivo, Herrería o Auxilio en Ruta", "3B82F6"),
        ("⬇", "", "64748B"),
        ("1. RECEPCIÓN", "Ruta/Auxilio Urgente vía WhatsApp  OR  Control Fluidos por Chequeador en Playa  OR  Preventivo", "1E3A8A"),
        ("⬇", "", "64748B"),
        ("2. DESPACHO & OT", "Parte de Taller marca Fuera de Servicio/Reparación/Auxilio + Supervisor crea OT en App Nativa", "1E3A8A"),
        ("⬇", "", "64748B"),
        ("3. ASIGNACIÓN", "Supervisor asigna Tareas (Preventivo/Correctivo/Herrería/Auxilio), Mecánico/Móvil e Insumos", "1E3A8A"),
        ("⬇", "", "64748B"),
        ("4. EJECUCIÓN", "Mecánico o Auxilio Móvil inicia Cronómetro en App Nativa, ejecuta trabajo/rescate y registra repuestos", "1E3A8A"),
        ("⬇", "", "64748B"),
        ("5. CONTROL", "Supervisor / Jefe de Taller realiza Inspección Final de Calidad y prueba de funcionamiento", "0F172A"),
        ("⬇", "", "64748B"),
        ("6. SYNC IA", "ASISTENTE IA (Sync Worker) sincroniza automáticamente con Taxes.com.ar y obtiene N° OT", "047857"),
        ("⬇", "", "64748B"),
        ("7. ARCHIVADO", "Auto-archivado a Historial + Exportación a Google Sheets (Mantenimiento y Auxilios) + Hugo AI", "047857"),
        ("⬇", "", "64748B"),
        ("FIN", "Unidad Marcada como ✅ OPERATIVA y lista para servicio en ruta", "166534")
    ]

    flow_table = doc.add_table(rows=len(flow_steps), cols=2)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, desc, color_hex) in enumerate(flow_steps):
        row = flow_table.rows[idx]
        cell_lbl = row.cells[0]
        cell_desc = row.cells[1]
        
        if label == "⬇":
            set_cell_background(cell_lbl, "FFFFFF")
            set_cell_background(cell_desc, "FFFFFF")
            set_cell_margins(cell_lbl, top=20, bottom=20, left=50, right=50)
            set_cell_margins(cell_desc, top=20, bottom=20, left=50, right=50)
            p = cell_lbl.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("⬇")
            r.font.size = Pt(11); r.bold = True; r.font.color.rgb = RGBColor(100, 116, 139)
            p2 = cell_desc.paragraphs[0]
            p2.paragraph_format.space_after = Pt(0)
        else:
            set_cell_background(cell_lbl, color_hex)
            set_cell_background(cell_desc, "F8FAFC")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=80, right=80)
            set_cell_margins(cell_desc, top=60, bottom=60, left=80, right=80)
            
            p = cell_lbl.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
            p2 = cell_desc.paragraphs[0]
            p2.paragraph_format.space_after = Pt(0)
            r2 = p2.add_run(desc)
            r2.font.name = 'Calibri'; r2.font.size = Pt(9.5)
            if label in ["INICIO", "FIN"]:
                r2.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 5. Descripción Etapa por Etapa
    add_heading_styled(doc, "5. Descripción Detallada del Flujo de Trabajo", 1)

    etapas = [
        ("Etapa 1: Inspección de Fluidos y Detección de Novedades", [
            "Control de Fluidos por Chequeador: El personal dedicado (Chequeador/Inspector) efectúa revisiones diarias en playa de niveles de aceite, líquido refrigerante, líquido de frenos e inspección visual general.",
            "Novedad o Deficiencia: Si el Chequeador detecta faltantes o anomalías, registra la novedad vía App Nativa o Chatbot.",
            "Auxilios en Ruta: Ante un desperfecto mecánico en servicio, el chofer notifica de inmediato la ubicación y detalles del auxilio vía Chatbot WhatsApp."
        ]),
        ("Etapa 2: Despacho de Auxilio y Creación de OT", [
            "El Supervisor activa el Protocolo de Auxilio Mecánico en Ruta y despacha la unidad móvil o mecánico de turno.",
            "Apertura de Orden de Trabajo (OT) en la App Nativa bajo clasificación: Preventivo, Correctivo, Herrería o Auxilio."
        ]),
        ("Etapa 3: Ejecución de Trabajos y Tiempos de Respuesta", [
            "El mecánico o auxilio móvil opera con tablet/celular desde el taller o lugar del rescate.",
            "Mide el tiempo real de reparación activando el cronómetro integrado en la App Nativa.",
            "Registra repuestos consumidos (correas, mangueras, cubiertas) y fluidos cargados."
        ]),
        ("Etapa 4: Control de Calidad y Retorno Operativo", [
            "El Supervisor / Jefe de Taller realiza la inspección de cierre del auxilio o reparación.",
            "Verifica que el 100% de las tareas estén en estado 'Finalizada'.",
            "Devuelve la unidad a estado '✅ Operativo' y aprueba el envío."
        ]),
        ("Etapa 5: Sincronización Automática Taxes.com.ar (ASISTENTE IA)", [
            "El ASISTENTE IA (Sync Worker Puppeteer) toma la orden de taller o auxilio en background.",
            "Accede a Taxes.com.ar, selecciona Rodado y Responsable sanitizado.",
            "Obtiene el número de OT oficial de Taxes y auto-archiva la orden a Historial."
        ]),
        ("Etapa 6: Análisis de Datos y Mejora Continua (ISO 9001)", [
            "Google Sheets 2026 actualiza indicadores KPI (MTBF, MTTR, frecuencia de auxilios por unidad, consumo de fluidos).",
            "Hugo AI Assistant disponible 24/7 en la app para responder dudas sobre auxilios pendientes, horas trabajadas y mecánicos en lenguaje natural."
        ])
    ]

    for title, items in etapas:
        add_heading_styled(doc, title, 2)
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run("▪ ")
            r1.font.color.rgb = RGBColor(37, 99, 235)
            r2 = p.add_run(item)
            r2.font.name = 'Calibri'; r2.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 6. Responsabilidades
    add_heading_styled(doc, "6. Matriz de Responsabilidades (RACI)", 1)
    
    resp_data = [
        ["Chofer / Conductor", "Notificar inmediatamente cualquier falla vía Chatbot WhatsApp antes o durante el viaje: novedades, auxilios, operatividad o deficiencia."],
        ["Chequeador / Inspector de Fluidos", "Realizar control diario de fluidos (aceite, agua, frenos) y estado general en playa; registrar novedades antes del despacho de la unidad."],
        ["Mecánico / Personal de Auxilio", "Acudir al auxilio mecánico en ruta, medir tiempos de reparación con cronómetro App Nativa, usar insumos correctamente y reportar desvíos."],
        ["Supervisor / Jefe de Taller", "Coordinar el despacho de auxilios, crear OTs (Preventivo/Correctivo/Herrería/Auxilio), auditar calidad del trabajo y validar el alta a Taxes.com.ar."],
        ["Asistente IA (Taxes - App Nativa)", "Garantizar la sincronización automática de OTs de taller y auxilios con Taxes.com.ar, respaldo de base de datos y actualización de Hugo AI."],
        ["Auditor Interno o Taxes - App Nativa IA", "Verificar el cumplimiento del procedimiento, seguimiento de KPI de auxilios e indicadores mensuales de falla."]
    ]

    resp_table = doc.add_table(rows=len(resp_data) + 1, cols=2)
    resp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    resp_headers = ["Rol / Puesto", "Responsabilidades Clave"]
    for c_idx, text in enumerate(resp_headers):
        cell = resp_table.rows[0].cells[c_idx]
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=90, bottom=90, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(9.5); r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row_data in enumerate(resp_data):
        row_cells = resp_table.rows[r_idx + 1].cells
        bg_hex = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Calibri'; r.font.size = Pt(9)
            if c_idx == 0:
                r.bold = True

    # Save output
    doc.save(output_path)
    
    # Save directly to public folder for direct browser download
    public_path = r"C:\Users\admin\.gemini\antigravity\scratch\app_taxes\public\Procedimiento_Mantenimiento_ISO9001.docx"
    doc.save(public_path)
    
    artifact_dir = r"C:\Users\admin\.gemini\antigravity\brain\d03ca953-b01d-40e0-af0b-75db698ee58a"
    if os.path.exists(artifact_dir):
        doc.save(os.path.join(artifact_dir, "Procedimiento_Mantenimiento_ISO9001.docx"))

    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    out = r"C:\Users\admin\.gemini\antigravity\scratch\app_taxes\Procedimiento_Mantenimiento_ISO9001_v3.2.docx"
    create_procedure_docx(out)
